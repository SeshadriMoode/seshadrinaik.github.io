"""Convert the JUM markdown manuscript to a Word file for Editorial Manager."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SRC = Path(__file__).resolve().parent / "Moode_JUM_manuscript.md"
OUT = Path(__file__).resolve().parent / "Moode_JUM_manuscript.docx"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITAL_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
CODE_RE = re.compile(r"`([^`]+)`")


def _set_run_font(run, size: int = 11, italic: bool = False, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold


def _add_runs(paragraph, text: str, size: int = 11) -> None:
    text = text.replace("\\[", "").replace("\\]", "").replace("\\(", "").replace("\\)", "")
    parts = BOLD_RE.split(text)
    for i, part in enumerate(parts):
        bold = i % 2 == 1
        sub = ITAL_RE.split(part)
        for j, chunk in enumerate(sub):
            italic = j % 2 == 1
            codes = CODE_RE.split(chunk)
            for k, bit in enumerate(codes):
                if not bit:
                    continue
                run = paragraph.add_run(bit)
                _set_run_font(run, size=size, italic=italic, bold=bold or (k % 2 == 1))


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(s) and s.startswith("|") and set(s.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set()


def _split_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def convert() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, line[2:].strip(), size=16)
            for run in p.runs:
                run.bold = True
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            _add_runs(p, line[3:].strip(), size=14)
            for run in p.runs:
                run.bold = True
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            _add_runs(p, line[4:].strip(), size=12)
            for run in p.runs:
                run.bold = True
                run.italic = True
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", line.strip()))
            i += 1
            continue
        if line.strip().startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            rows = [_split_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    table.rows[r].cells[c].text = ""
                    para = table.rows[r].cells[c].paragraphs[0]
                    _add_runs(para, cell.replace("**", ""), size=9)
                    if r == 0:
                        for run in para.runs:
                            run.bold = True
            doc.add_paragraph()
            continue
        if line.strip() == "":
            i += 1
            continue
        p = doc.add_paragraph()
        _add_runs(p, line.strip())
        i += 1

    # Embed figures at the end in reading order
    fig_dir = Path(__file__).resolve().parent / "figures"
    captions = {
        "Fig1_study_area.png": "Fig. 1. Eixos Verds phase 1 on an Esri street map of central Barcelona, with as-built green axes and stations by network distance.",
        "Fig2_event_study.png": "Fig. 2. Event-study DiD versus 2019, working-day MADT versus the >2 km control.",
        "Fig3_window_comparison.png": "Fig. 3. Nearby volume change under cumulative, drop-2020 and incremental windows.",
        "Fig4_incremental_did.png": "Fig. 4. Incremental DiD by network-distance bin, 2022 H1 versus 2024–25.",
        "Fig5_ltr_nr_nvtr.png": "Fig. 5. Simulated demand-fixed LTR, NR and NVTR (not measured network VKT).",
        "Fig6_dvkt_by_ring.png": "Fig. 6. Simulated demand-fixed ΔVKT by network ring.",
        "Fig7_observed_vs_simulated.png": "Fig. 7. Observed incremental ΔQ versus simulated demand-fixed ΔQ.",
    }
    p = doc.add_paragraph()
    _add_runs(p, "Figures", size=14)
    for run in p.runs:
        run.bold = True
    for name, cap in captions.items():
        path = fig_dir / name
        if not path.exists():
            continue
        doc.add_picture(str(path), width=Inches(6.2))
        cap_p = doc.add_paragraph()
        _add_runs(cap_p, cap, size=10)
        for run in cap_p.runs:
            run.italic = True

    doc.save(OUT)
    print("wrote", OUT, "bytes", OUT.stat().st_size)


if __name__ == "__main__":
    convert()
